import os
import re
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import hashlib
from datetime import datetime

class EnhancedWordLoader:
    """增强版Word文档加载器，支持语义感知切分和多粒度索引"""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
        )
    
    def extract_document_structure(self, doc: DocumentType) -> Dict[str, Any]:
        """提取文档结构信息"""
        structure = {
            'title': '',
            'sections': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        # 提取标题（假设第一个非空段落是标题）
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                structure['title'] = paragraph.text.strip()
                break
        
        # 提取章节信息
        current_section = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # 检测章节标题（基于格式或关键词）
            if self._is_section_header(paragraph):
                current_section = {
                    'title': text,
                    'level': self._get_header_level(paragraph),
                    'content': []
                }
                structure['sections'].append(current_section)
            elif current_section:
                current_section['content'].append(text)
        
        return structure
    
    def _is_section_header(self, paragraph: Paragraph) -> bool:
        """判断段落是否为章节标题"""
        text = paragraph.text.strip()
        
        # 检查是否包含数字编号
        if re.match(r'^[\d\.]+\s+', text):
            return True
        
        # 检查是否包含常见标题关键词
        header_keywords = ['章', '节', '第', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if any(keyword in text for keyword in header_keywords):
            return True
        
        # 检查段落格式（字体大小、加粗等）
        if paragraph.style.name.startswith('Heading'):
            return True
        
        return False
    
    def _get_header_level(self, paragraph: Paragraph) -> int:
        """获取标题级别"""
        style_name = paragraph.style.name.lower()
        if 'heading 1' in style_name or 'title' in style_name:
            return 1
        elif 'heading 2' in style_name:
            return 2
        elif 'heading 3' in style_name:
            return 3
        else:
            return 1
    
    def semantic_chunking(self, text: str, structure: Dict[str, Any]) -> List[Dict[str, Any]]:
        """基于语义的文本切分"""
        chunks = []
        
        # 按章节切分
        for section in structure['sections']:
            section_text = '\n'.join(section['content'])
            if len(section_text) <= self.chunk_size:
                # 短章节，作为一个块
                chunks.append({
                    'content': section_text,
                    'type': 'section',
                    'title': section['title'],
                    'level': section['level'],
                    'metadata': {
                        'section_title': section['title'],
                        'section_level': section['level'],
                        'chunk_type': 'section'
                    }
                })
            else:
                # 长章节，进一步切分
                sub_chunks = self.text_splitter.split_text(section_text)
                for i, chunk in enumerate(sub_chunks):
                    chunks.append({
                        'content': chunk,
                        'type': 'subsection',
                        'title': f"{section['title']} - 第{i+1}部分",
                        'level': section['level'] + 1,
                        'metadata': {
                            'section_title': section['title'],
                            'section_level': section['level'],
                            'chunk_type': 'subsection',
                            'chunk_index': i
                        }
                    })
        
        return chunks
    
    def multi_granularity_chunking(self, text: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """多粒度切分：长块（段落级）和短块（句子级）"""
        
        # 长块切分（段落级）
        long_chunks = self.text_splitter.split_text(text)
        long_chunk_docs = []
        for i, chunk in enumerate(long_chunks):
            long_chunk_docs.append({
                'content': chunk,
                'type': 'long_chunk',
                'metadata': {
                    'chunk_type': 'long_chunk',
                    'chunk_index': i,
                    'granularity': 'paragraph'
                }
            })
        
        # 短块切分（句子级）
        sentences = re.split(r'[。！？.!?]', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        short_chunk_docs = []
        for i, sentence in enumerate(sentences):
            if len(sentence) > 10:  # 过滤太短的句子
                short_chunk_docs.append({
                    'content': sentence,
                    'type': 'short_chunk',
                    'metadata': {
                        'chunk_type': 'short_chunk',
                        'chunk_index': i,
                        'granularity': 'sentence'
                    }
                })
        
        return long_chunk_docs, short_chunk_docs
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """加载Word文档并返回多粒度切分结果"""
        try:
            doc = Document(file_path)
            
            # 提取文档结构
            structure = self.extract_document_structure(doc)
            
            # 提取所有文本
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # 生成文档ID
            doc_id = hashlib.md5(f"{file_path}_{datetime.now().isoformat()}".encode()).hexdigest()
            
            # 基础元数据
            base_metadata = {
                'source': file_path,
                'doc_id': doc_id,
                'title': structure['title'],
                'upload_time': datetime.now().isoformat(),
                'file_size': os.path.getsize(file_path),
                'total_sections': len(structure['sections'])
            }
            
            # 语义切分
            semantic_chunks = self.semantic_chunking(full_text, structure)
            
            # 多粒度切分
            long_chunks, short_chunks = self.multi_granularity_chunking(full_text)
            
            # 合并所有切分结果
            all_chunks = []
            
            # 添加语义切分结果
            for chunk in semantic_chunks:
                chunk_metadata = {**base_metadata, **chunk['metadata']}
                all_chunks.append(LangchainDocument(
                    page_content=chunk['content'],
                    metadata=chunk_metadata
                ))
            
            # 添加长块切分结果
            for chunk in long_chunks:
                chunk_metadata = {**base_metadata, **chunk['metadata']}
                all_chunks.append(LangchainDocument(
                    page_content=chunk['content'],
                    metadata=chunk_metadata
                ))
            
            # 添加短块切分结果
            for chunk in short_chunks:
                chunk_metadata = {**base_metadata, **chunk['metadata']}
                all_chunks.append(LangchainDocument(
                    page_content=chunk['content'],
                    metadata=chunk_metadata
                ))
            
            return {
                'documents': all_chunks,
                'structure': structure,
                'metadata': base_metadata,
                'statistics': {
                    'total_chunks': len(all_chunks),
                    'semantic_chunks': len(semantic_chunks),
                    'long_chunks': len(long_chunks),
                    'short_chunks': len(short_chunks)
                }
            }
            
        except Exception as e:
            raise ValueError(f"处理文档失败: {e}")
    
    def load_directory(self, directory_path: str) -> Dict[str, Any]:
        """加载目录中的所有Word文档"""
        word_files = [f for f in os.listdir(directory_path) if f.endswith('.docx')]
        all_docs = []
        all_structures = []
        all_metadata = []
        
        for word_file in word_files:
            file_path = os.path.join(directory_path, word_file)
            result = self.load(file_path)
            all_docs.extend(result['documents'])
            all_structures.append(result['structure'])
            all_metadata.append(result['metadata'])
        
        return {
            'documents': all_docs,
            'structures': all_structures,
            'metadata': all_metadata,
            'total_files': len(word_files)
        }
